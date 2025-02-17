import streamlit as st
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re

# Función para limpiar Markdown
def clean_markdown(text):
    """Elimina marcas de Markdown del texto."""
    text = re.sub(r'[#*_`]', '', text)  # Eliminar caracteres especiales de Markdown
    return text.strip()

# Función para aplicar reglas de capitalización según el idioma
def format_title(title, language):
    """
    Formatea el título según las reglas gramaticales del idioma.
    - Español: Solo mayúscula inicial en la primera palabra y nombres propios.
    - Otros idiomas: Mayúscula inicial en cada palabra.
    """
    if language.lower() == "spanish":
        # Dividir el título en palabras
        words = title.split()
        # Mantener mayúscula inicial solo en la primera palabra y nombres propios
        formatted_words = [words[0].capitalize()] + [word.lower() for word in words[1:]]
        return " ".join(formatted_words)
    else:
        # Capitalizar cada palabra para otros idiomas
        return title.title()

# Función para generar un capítulo
def generate_chapter(api_key, topic, audience, chapter_number, language, instructions="", is_intro=False, is_conclusion=False):
    url = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    # Construir el mensaje con las instrucciones especiales
    if is_intro:
        message_content = f"Write the introduction of a book about {topic} aimed at {audience} with 500-800 words in {language}."
    elif is_conclusion:
        message_content = f"Write the conclusions of a book about {topic} aimed at {audience} with 500-800 words in {language}."
    else:
        message_content = f"Write chapter {chapter_number} of a book about {topic} aimed at {audience} with 2000-2500 words in {language}."
    
    if instructions:
        message_content += f" Additional instructions: {instructions}"
    
    data = {
        "model": "qwen-turbo",
        "messages": [
            {"role": "system", "content": f"You are a helpful assistant that writes in {language}."},
            {"role": "user", "content": message_content}
        ]
    }
    try:
        response = requests.post(url, json=data, headers=headers)
        response.raise_for_status()  # Lanza una excepción si hay un error HTTP
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "Error generating the chapter.")
    except Exception as e:
        st.error(f"Error generating chapter {chapter_number}: {str(e)}")
        content = "Error generating the chapter."
    return clean_markdown(content)

# Función para agregar numeración de páginas al documento Word
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)

# Función para crear un documento Word con formato específico
def create_word_document(chapters, title, author_name, author_bio, language):
    doc = Document()

    # Configurar el tamaño de página (5.5 x 8.5 pulgadas)
    section = doc.sections[0]
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)

    # Configurar márgenes de 0.8 pulgadas en todo
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Añadir título formateado según el idioma
    formatted_title = format_title(title, language)
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    title_run = title_paragraph.add_run(formatted_title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"

    # Añadir nombre del autor si está proporcionado
    if author_name:
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
        author_run = author_paragraph.add_run(author_name)
        author_run.font.size = Pt(12)
        author_run.font.name = "Times New Roman"
        doc.add_page_break()  # Salto de página después del título y autor

    # Añadir perfil del autor si está proporcionado
    if author_bio:
        bio_paragraph = doc.add_paragraph("Author Information")
        bio_paragraph.style = "Heading 2"
        bio_paragraph.runs[0].font.size = Pt(11)
        bio_paragraph.runs[0].font.name = "Times New Roman"
        doc.add_paragraph(author_bio).style = "Normal"
        doc.add_page_break()  # Salto de página después del perfil del autor

    # Añadir capítulos
    for i, chapter in enumerate(chapters, 1):
        # Añadir encabezado del capítulo formateado según el idioma
        chapter_title_text = f"Chapter {i}" if language.lower() != "spanish" else f"Capítulo {i}"
        formatted_chapter_title = format_title(chapter_title_text, language)
        chapter_title = doc.add_paragraph(formatted_chapter_title)
        chapter_title.style = "Heading 1"
        chapter_title.runs[0].font.size = Pt(12)
        chapter_title.runs[0].font.name = "Times New Roman"

        # Insertar el contenido del capítulo como bloques de párrafos
        paragraphs = chapter.split('\n\n')  # Dividir por párrafos (asumiendo doble espacio entre párrafos)
        for para_text in paragraphs:
            paragraph = doc.add_paragraph(para_text.strip())  # Crear un nuevo párrafo
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada
            paragraph.paragraph_format.space_after = Pt(6)  # Espaciado posterior entre párrafos
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        doc.add_page_break()  # Salto de página entre capítulos

    # Agregar numeración de páginas
    add_page_numbers(doc)

    # Guardar el documento en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Configuración de Streamlit
st.set_page_config(
    page_title="Automatic Book Generator",
    page_icon="📚",  # Ícono para la pestaña del navegador
)

# Título con ícono
st.title("📚 Automatic Book Generator")

# Barra lateral con instrucciones y anuncio
st.sidebar.header("📖 How does this app work?")
st.sidebar.markdown("""
This application automatically generates non-fiction books in `.docx` format based on a topic and target audience.  
**Steps to use it:**
1. Enter the book's topic.
2. Specify the target audience.
3. Write special instructions (optional).
4. Select the number of chapters desired (maximum 20).
5. Choose the book's language.
6. Decide whether to include an introduction, conclusions, author name, and author profile.
7. Click "Generate Book".
8. Download the generated file.
""")
st.sidebar.markdown("""
---
**📝 Text correction in 24 hours**  
👉 [Hablemos Bien](https://hablemosbien.org)
""")

# Validación de claves secretas
if "DASHSCOPE_API_KEY" not in st.secrets:
    st.error("Please configure the API key in Streamlit secrets.")
    st.stop()
api_key = st.secrets["DASHSCOPE_API_KEY"]

# Entradas del usuario
topic = st.text_input("📒 Book Topic:")
audience = st.text_input("🎯 Target Audience:")
instructions = st.text_area("📝 Special Instructions (optional):", 
                             placeholder="Example: Use a formal tone, include practical examples, avoid technical jargon...")
num_chapters = st.slider("🔢 Number of Chapters", min_value=1, max_value=20, value=5)

# Opciones para introducción y conclusiones
include_intro = st.checkbox("Include Introduction", value=True)
include_conclusion = st.checkbox("Include Conclusions", value=True)

# Opciones adicionales
author_name = st.text_input("🖋️ Author Name (optional):")
author_bio = st.text_area("👤 Author Profile (optional):", 
                          placeholder="Example: Brief professional description or biography.")

# Menú desplegable para elegir el idioma
languages = [
    "English", "Spanish", "French", "German", "Chinese", "Japanese", 
    "Russian", "Portuguese", "Italian", "Arabic", "Medieval Latin", "Koine Greek"
]
selected_language = st.selectbox("🌐 Choose the book's language:", languages)

# Estado de Streamlit para almacenar los capítulos generados
if 'chapters' not in st.session_state:
    st.session_state.chapters = []

# Botón para generar el libro
if st.button("🚀 Generate Book"):
    if not topic or not audience:
        st.error("Please enter a valid topic and target audience.")
        st.stop()
     
    chapters = []
    
    # Generar introducción si está seleccionada
    if include_intro:
        st.write("⏳ Generating introduction...")
        intro_content = generate_chapter(api_key, topic, audience, 0, selected_language.lower(), instructions, is_intro=True)
        chapters.append(intro_content)
        with st.expander("🌟 Introduction"):
            st.write(intro_content)
    
    # Generar capítulos principales
    progress_bar = st.progress(0)
    for i in range(1, num_chapters + 1):
        st.write(f"⏳ Generating chapter {i}...")
        chapter_content = generate_chapter(api_key, topic, audience, i, selected_language.lower(), instructions)
        word_count = len(chapter_content.split())   # Contar palabras
        chapters.append(chapter_content)
        with st.expander(f" Chapter {i} ({word_count} words)"):
            st.write(chapter_content)
        progress_bar.progress(i / num_chapters)
    
    # Generar conclusiones si están seleccionadas
    if include_conclusion:
        st.write("⏳ Generating conclusions...")
        conclusion_content = generate_chapter(api_key, topic, audience, 0, selected_language.lower(), instructions, is_conclusion=True)
        chapters.append(conclusion_content)
        with st.expander("🔚 Conclusions"):
            st.write(conclusion_content)
    
    # Almacenar los capítulos en el estado de Streamlit
    st.session_state.chapters = chapters

# Mostrar opciones de descarga si hay capítulos generados
if st.session_state.chapters:
    st.subheader("⬇️ Download Options")
    word_file = create_word_document(st.session_state.chapters, topic, author_name, author_bio, selected_language.lower())

    st.download_button(
        label="📥 Download in Word",
        data=word_file,
        file_name=f"{topic}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
